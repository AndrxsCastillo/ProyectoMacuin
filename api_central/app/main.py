from fastapi import FastAPI, Depends, HTTPException, status, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from sqlalchemy.orm import Session
from typing import List
from . import models, schemas, crud
from .database import SessionLocal, engine
from datetime import datetime, timezone
from sqlalchemy import func
from fastapi.responses import StreamingResponse
import io
from fpdf import FPDF
import openpyxl
from docx import Document

# (Opcional) Crea las tablas en MySQL si no existen, aunque ya corrimos el script SQL
models.Base.metadata.create_all(bind=engine)

app = FastAPI(
    title="API Central - MACUIN",
    description="Núcleo del sistema para gestión de autopartes y usuarios.",
    version="1.0.0"
)

# ==========================================
# CONFIGURACIÓN CORS (Permite que Laravel y Flask se conecten)
# ==========================================
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción se cambian por las URLs de Laravel y Flask
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Dependencia para la base de datos
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.get("/", tags=["Inicio"])
async def inicio():
    return  {"mensaje": "Bienvenido a la API del proyecto Macuin :)"}

# ==========================================
# ENDPOINT DE LOGIN (Generación de Token)
# ==========================================
@app.post("/token", response_model=schemas.Token, tags=["Seguridad"])
def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    usuario = crud.get_usuario_by_email(db, email=form_data.username)
    # Aquí deberíamos verificar el hash de la contraseña. 
    # Por simplicidad en este paso, solo verificamos que el correo exista.
    if not usuario:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Correo o contraseña incorrectos",
            headers={"WWW-Authenticate": "Bearer"},
        )
    # En un sistema real aquí generamos el JWT cifrado
    return {"access_token": usuario.email, "token_type": "bearer"}

# ==========================================
# RUTAS DE USUARIOS (M-11, M-12, M-13, M-14)
# ==========================================
@app.post("/usuarios/", response_model=schemas.Usuario, tags=["Usuarios"])
def crear_usuario(usuario: schemas.UsuarioCreate, db: Session = Depends(get_db)):
    db_usuario = crud.get_usuario_by_email(db, email=usuario.email)
    if db_usuario:
        raise HTTPException(status_code=400, detail="Este email ya está registrado")
    return crud.create_usuario(db=db, usuario=usuario)

@app.get("/usuarios/", tags=["Usuarios"])
def leer_usuarios(db: Session = Depends(get_db)):
    usuarios_db = db.query(models.Usuario).all()
    lista_usuarios = []
    
    for u in usuarios_db:
        rol_db = db.query(models.Rol).filter(models.Rol.id == u.rol_id).first()
        nombre_rol = rol_db.nombre if rol_db else "Sin Rol"
        
        lista_usuarios.append({
            "id": u.id,
            "nombre": u.nombre,
            "email": u.email,
            "activo": u.activo,
            "rol": nombre_rol  # Esto es lo que Flask necesita leer
        })
        
    return lista_usuarios

@app.put("/usuarios/{usuario_id}", response_model=schemas.Usuario, tags=["Usuarios"])
def actualizar_usuario(usuario_id: int, usuario: schemas.UsuarioUpdate, db: Session = Depends(get_db)):
    db_usuario = crud.update_usuario(db, usuario_id=usuario_id, usuario_data=usuario)
    if db_usuario is None:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")
    return db_usuario

@app.delete("/usuarios/{usuario_id}", response_model=schemas.Usuario, tags=["Usuarios"])
def eliminar_usuario(usuario_id: int, db: Session = Depends(get_db)):
    db_usuario = crud.delete_usuario(db, usuario_id=usuario_id)
    if db_usuario is None:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")
    return db_usuario

# ==========================================
# RUTAS DE CATEGORÍAS
# ==========================================
@app.get("/categorias/", response_model=List[schemas.Categoria], tags=["Categorías"])
def leer_categorias(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    categorias = crud.get_categorias(db, skip=skip, limit=limit)
    return categorias

# ==========================================
# RUTAS DE AUTOPARTES (M-07, M-15, M-16)
# ==========================================
@app.post("/autopartes/", response_model=schemas.Autoparte, tags=["Autopartes"])
def crear_autoparte(autoparte: schemas.AutoparteCreate, db: Session = Depends(get_db)):
    return crud.create_autoparte(db=db, autoparte=autoparte)

@app.get("/autopartes/", response_model=List[schemas.Autoparte], tags=["Autopartes"])
def leer_autopartes(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    autopartes = crud.get_autopartes(db, skip=skip, limit=limit)
    return autopartes

@app.get("/autopartes/{autoparte_id}", tags=["Autopartes"])
def leer_autoparte(autoparte_id: int, db: Session = Depends(get_db)):
    # 1. Traemos la info de la autoparte
    db_autoparte = crud.get_autoparte(db, autoparte_id=autoparte_id)
    if db_autoparte is None:
        raise HTTPException(status_code=404, detail="Autoparte no encontrada")
        
    # 2. Traemos su inventario de la tabla separada
    db_inventario = db.query(models.Inventario).filter(models.Inventario.autoparte_id == autoparte_id).first()
    
    # 3. Armamos un diccionario "libre" con todo junto
    return {
        "id": db_autoparte.id,
        "nombre": db_autoparte.nombre,
        "descripcion": db_autoparte.descripcion,
        "categoria_id": db_autoparte.categoria_id,
        "marca": db_autoparte.marca,
        "precio": float(db_autoparte.precio),
        "activo": db_autoparte.activo,
        # Si encuentra inventario manda el real, si no, manda 0
        "stock_actual": db_inventario.stock_actual if db_inventario else 0,
        "stock_minimo": db_inventario.stock_minimo if db_inventario else 0
    }

@app.put("/autopartes/{autoparte_id}", response_model=schemas.Autoparte, tags=["Autopartes"])
def actualizar_autoparte(autoparte_id: int, autoparte: schemas.AutoparteUpdate, db: Session = Depends(get_db)):
    db_autoparte = crud.update_autoparte(db, autoparte_id=autoparte_id, autoparte_data=autoparte)
    if db_autoparte is None:
        raise HTTPException(status_code=404, detail="Autoparte no encontrada")
    return db_autoparte

@app.delete("/autopartes/{autoparte_id}", response_model=schemas.Autoparte, tags=["Autopartes"])
def eliminar_autoparte(autoparte_id: int, db: Session = Depends(get_db)):
    db_autoparte = crud.delete_autoparte(db, autoparte_id=autoparte_id)
    if db_autoparte is None:
        raise HTTPException(status_code=404, detail="Autoparte no encontrada")
    return db_autoparte

# ==========================================
# RUTAS DE PEDIDOS
# ==========================================
@app.get("/pedidos/", response_model=List[schemas.PedidoBase], tags=["Pedidos"])
def leer_pedidos(db: Session = Depends(get_db)):
    # Usamos JOIN para mezclar la tabla pedidos con estados_pedido
    # Así Flask recibe la palabra (ej. "CREADO") y no solo el número (ej. 1)
    resultados = db.query(
        models.Pedido.id,
        models.Pedido.fecha_pedido,
        models.Pedido.total,
        models.EstadoPedido.nombre.label("estatus")
    ).join(models.EstadoPedido, models.Pedido.estado_id == models.EstadoPedido.id).all()
    
    return resultados

from fastapi import HTTPException

@app.put("/pedidos/{id}/estatus", tags=["Pedidos"])
def actualizar_estatus_pedido(id: int, estatus_data: schemas.PedidoActualizarEstatus, db: Session = Depends(get_db)):
    pedido = db.query(models.Pedido).filter(models.Pedido.id == id).first()
    if not pedido:
        raise HTTPException(status_code=404, detail="Pedido no encontrado")
    
    nuevo_estado = db.query(models.EstadoPedido).filter(models.EstadoPedido.nombre == estatus_data.estatus).first()
    if not nuevo_estado:
        raise HTTPException(status_code=400, detail="El estado proporcionado no existe en el catálogo")

    estado_actual = db.query(models.EstadoPedido).filter(models.EstadoPedido.id == pedido.estado_id).first()

    # LÓGICA DE INVENTARIO CORREGIDA: Usamos .upper() para evitar fallos de texto
    if nuevo_estado.nombre.upper() == 'SURTIDO' and estado_actual.nombre.upper() != 'SURTIDO':
        for detalle in pedido.detalles:
            inventario = db.query(models.Inventario).filter(models.Inventario.autoparte_id == detalle.autoparte_id).first()
            
            if not inventario:
                raise HTTPException(status_code=400, detail=f"No hay registro de inventario para la pieza ID {detalle.autoparte_id}")
            
            if inventario.stock_actual < detalle.cantidad:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Stock insuficiente para la pieza ID {detalle.autoparte_id}. Solicitado: {detalle.cantidad}, Disponible: {inventario.stock_actual}"
                )
            
            # Si hay stock, lo descontamos
            inventario.stock_actual -= detalle.cantidad

            # NUEVO METODO: Usamos now con la zona horaria UTC explícita
            inventario.fecha_actualizacion = datetime.now(timezone.utc)

    # Actualizamos el estado del pedido
    pedido.estado_id = nuevo_estado.id
    db.commit()
    
    return {"mensaje": "Estatus actualizado correctamente", "estatus": nuevo_estado.nombre}

# ==========================================
# RUTAS DE REPORTES
# ==========================================
@app.get("/reportes/", tags=["Reportes"])
def obtener_reportes(db: Session = Depends(get_db)):
    # 1. Ingresos Totales
    ingresos_totales = db.query(func.sum(models.Pedido.total)).scalar() or 0.0

    # 2. Resumen de Pedidos por Estatus
    pedidos_por_estatus = db.query(
        models.EstadoPedido.nombre.label("estatus"),
        func.count(models.Pedido.id).label("cantidad")
    ).outerjoin(models.Pedido, models.EstadoPedido.id == models.Pedido.estado_id)\
     .group_by(models.EstadoPedido.nombre).all()
    resumen_pedidos = {item.estatus: item.cantidad for item in pedidos_por_estatus}

    # 3. Alertas de Inventario Crítico
    inventario_bajo = db.query(
        models.Autoparte.nombre, models.Inventario.stock_actual, models.Inventario.stock_minimo
    ).join(models.Autoparte).filter(models.Inventario.stock_actual <= models.Inventario.stock_minimo).limit(5).all()
    alertas_stock = [{"pieza": i.nombre, "stock_actual": i.stock_actual, "stock_minimo": i.stock_minimo} for i in inventario_bajo]

    # 4. Reporte de Clientes (Top 5 por volumen de compra)
    top_clientes = db.query(
        models.Usuario.nombre,
        models.Usuario.email,
        func.count(models.Pedido.id).label("total_pedidos"),
        func.sum(models.Pedido.total).label("total_comprado")
    ).join(models.Pedido).group_by(models.Usuario.id).order_by(func.sum(models.Pedido.total).desc()).limit(5).all()
    reporte_clientes = [{"nombre": c.nombre, "email": c.email, "pedidos": c.total_pedidos, "inversion": float(c.total_comprado)} for c in top_clientes]

    # 5. Reporte de Ventas (Top 5 productos más vendidos)
    top_ventas = db.query(
        models.Autoparte.nombre,
        func.sum(models.DetallePedido.cantidad).label("piezas_vendidas"),
        func.sum(models.DetallePedido.cantidad * models.DetallePedido.precio_unitario).label("ingreso_generado")
    ).join(models.DetallePedido).group_by(models.Autoparte.id).order_by(func.sum(models.DetallePedido.cantidad).desc()).limit(5).all()
    reporte_ventas = [{"pieza": v.nombre, "vendidas": v.piezas_vendidas, "ingreso": float(v.ingreso_generado)} for v in top_ventas]

    return {
        "ingresos_totales": float(ingresos_totales),
        "resumen_pedidos": resumen_pedidos,
        "alertas_stock": alertas_stock,
        "reporte_clientes": reporte_clientes,
        "reporte_ventas": reporte_ventas
    }

# ==========================================
# RUTAS DE ROLES
# ==========================================
@app.get("/roles/", tags=["Usuarios"])
def leer_roles(db: Session = Depends(get_db)):
    return db.query(models.Rol).all()

# ==========================================
# MOTOR GENERADOR DE REPORTES (PDF, XLSX, DOCX)
# ==========================================
def generar_archivo(formato: str, titulo: str, columnas: list, datos: list):
    """Motor genérico para exportar a PDF, Excel o Word"""
    
    if formato == "pdf":
        pdf = FPDF(orientation='L') # Orientación Horizontal (Landscape) para que quepan las tablas
        pdf.add_page()
        pdf.set_font("helvetica", "B", 16)
        
        # Color azul corporativo para el título
        pdf.set_text_color(30, 58, 138) # #1e3a8a
        pdf.cell(0, 10, titulo, ln=True, align="C")
        pdf.ln(10)
        
        # Color y fuente para encabezados de tabla
        pdf.set_font("helvetica", "B", 10)
        pdf.set_fill_color(30, 58, 138)
        pdf.set_text_color(255, 255, 255)
        
        ancho_col = 270 / len(columnas) # 270 es aprox el ancho utilizable en Landscape
        for col in columnas:
            pdf.cell(ancho_col, 10, str(col), border=1, align="C", fill=True)
        pdf.ln()
        
        # Filas de datos
        pdf.set_font("helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        for fila in datos:
            for item in fila:
                # Recortamos el texto a 30 caracteres para que no se desborde la celda
                texto = str(item)
                if len(texto) > 30:
                    texto = texto[:27] + "..."
                pdf.cell(ancho_col, 10, texto, border=1, align="C")
            pdf.ln()
            
        pdf_bytes = pdf.output()
        return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf")

    elif formato == "xlsx":
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Reporte MACUIN"
        
        # Encabezados
        ws.append(columnas)
        # Datos
        for fila in datos:
            ws.append(fila)
        
        # Auto-ajustar ancho de columnas básico
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter 
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            ws.column_dimensions[column].width = adjusted_width

        stream = io.BytesIO()
        wb.save(stream)
        stream.seek(0)
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    elif formato == "docx":
        doc = Document()
        doc.add_heading(titulo, 0)
        
        tabla = doc.add_table(rows=1, cols=len(columnas))
        tabla.style = 'Table Grid'
        
        # Encabezados
        hdr_cells = tabla.rows[0].cells
        for i, col in enumerate(columnas):
            hdr_cells[i].text = str(col)
            
        # Datos
        for fila in datos:
            row_cells = tabla.add_row().cells
            for i, item in enumerate(fila):
                row_cells[i].text = str(item)
                
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    raise HTTPException(status_code=400, detail="Formato no soportado")


@app.get("/reportes/generar/{tipo}", tags=["Reportes"])
def descargar_reporte(
    tipo: str, 
    formato: str = Query("pdf"),
    categoria: str = Query(None),
    rol: str = Query(None),
    fecha_inicio: str = Query(None),
    fecha_fin: str = Query(None),
    db: Session = Depends(get_db)
):
    titulo = ""
    columnas = []
    datos_matriz = []

    # 1. REPORTE DE INVENTARIO (Usa JOIN con la tabla Categoria e Inventario)
    if tipo == "inventario":
        titulo = "Catálogo de Inventario MACUIN"
        columnas = ["ID", "Nombre", "Marca", "Categoría", "Stock Actual", "Precio"]
        
        # Hacemos la consulta base uniendo Autoparte, Categoria e Inventario
        query = db.query(
            models.Autoparte.id,
            models.Autoparte.nombre,
            models.Autoparte.marca,
            models.Categoria.nombre.label('categoria_nombre'),
            models.Inventario.stock_actual,
            models.Autoparte.precio
        ).join(models.Categoria).outerjoin(models.Inventario)

        piezas = query.all()
        
        for p in piezas:
            # Si el filtro de categoría está activo, omitimos las que no coincidan
            if categoria and categoria != "todas":
                # Convertimos todo a minúsculas y comparamos sin espacios
                cat_db = str(p.categoria_nombre).lower().replace(" ", "")
                cat_filtro = str(categoria).lower()
                if cat_filtro not in cat_db:
                    continue

            stock = p.stock_actual if p.stock_actual is not None else 0
            datos_matriz.append([p.id, p.nombre, p.marca, p.categoria_nombre, stock, f"${p.precio}"])

    # 2. REPORTE DE USUARIOS (Usa JOIN con la tabla Rol)
    elif tipo == "usuarios":
        titulo = "Directorio de Usuarios del Sistema"
        columnas = ["ID", "Nombre", "Email", "Rol", "Estado"]
        
        query = db.query(
            models.Usuario.id,
            models.Usuario.nombre,
            models.Usuario.email,
            models.Rol.nombre.label('rol_nombre'),
            models.Usuario.activo
        ).join(models.Rol)

        usuarios = query.all()
        
        for u in usuarios:
            if rol and rol != "todos":
                rol_db = str(u.rol_nombre).lower()
                if rol != rol_db:
                    continue

            estado = "Activo" if u.activo else "Inactivo"
            datos_matriz.append([u.id, u.nombre, u.email, u.rol_nombre, estado])

    # 3. REPORTE DE ALERTAS (Stock Crítico)
    elif tipo == "alertas":
        titulo = "Reporte de Stock Crítico (Urgente)"
        columnas = ["Pieza", "Marca", "Stock Actual", "Mínimo Permitido"]
        
        # Filtramos donde stock_actual es menor o igual al stock_minimo
        piezas_criticas = db.query(
            models.Autoparte.nombre, 
            models.Autoparte.marca, 
            models.Inventario.stock_actual, 
            models.Inventario.stock_minimo
        ).join(models.Inventario).filter(models.Inventario.stock_actual <= models.Inventario.stock_minimo).all()
        
        for p in piezas_criticas:
            datos_matriz.append([p.nombre, p.marca, p.stock_actual, p.stock_minimo])

    # 4. REPORTE DE VENTAS (Con filtro de fechas)
    elif tipo == "ventas":
        titulo = f"Historial de Ventas ({fecha_inicio} a {fecha_fin})"
        columnas = ["ID Pedido", "Fecha", "Estatus", "Total"]
        
        query = db.query(
            models.Pedido.id,
            models.Pedido.fecha_pedido,
            models.EstadoPedido.nombre.label('estatus'),
            models.Pedido.total
        ).join(models.EstadoPedido)

        # TODO: Aquí podrías agregar el .filter() de fechas si tu BD guarda las fechas en formato compatible.
        # Por ahora lo traemos todo para evitar errores de parseo de fechas en MySQL.
        pedidos = query.all()
        
        for p in pedidos:
            # Formateamos la fecha si existe
            fecha_str = p.fecha_pedido.strftime("%Y-%m-%d") if p.fecha_pedido else "N/A"
            datos_matriz.append([p.id, fecha_str, p.estatus, f"${p.total}"])

    else:
        raise HTTPException(status_code=404, detail="Tipo de reporte no existe")

    # Si después de filtrar no hay datos, agregamos una fila indicándolo para que no truene el archivo
    if not datos_matriz:
        datos_matriz.append(["-", "No hay datos", "para estos", "filtros", "-", "-"])

    return generar_archivo(formato, titulo, columnas, datos_matriz)