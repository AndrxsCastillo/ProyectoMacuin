import io
from typing import List, Iterable
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from fpdf import FPDF
import openpyxl
from docx import Document


def _normalize_rows(columnas: List[str], datos: Iterable) -> List[List[str]]:
    """Normaliza las filas: convierte a listas, rellena/trunca según columnas."""
    cols_len = len(columnas) if columnas else 0
    normalized = []
    for row in datos or []:
        if not isinstance(row, (list, tuple)):
            row = [row]
        row_list = ["" if v is None else v for v in row]
        if cols_len:
            if len(row_list) < cols_len:
                row_list = row_list + [""] * (cols_len - len(row_list))
            elif len(row_list) > cols_len:
                row_list = row_list[:cols_len]
        normalized.append(row_list)
    return normalized


def generar_archivo(formato: str, titulo: str, columnas: List[str], datos: Iterable):
    """Genera y devuelve un StreamingResponse con el archivo en el formato pedido.

    Valid formats: 'pdf', 'xlsx', 'docx'.
    - `columnas` debe ser lista (puede estar vacía).
    - `datos` debe ser iterable de filas (cada fila: lista/tupla).
    La función normaliza filas para que coincidan con el número de columnas (si se
    proveen columnas), y realiza validaciones para devolver errores HTTP claros.
    """

    if not isinstance(formato, str) or not formato:
        raise HTTPException(status_code=400, detail="Formato inválido")
    fmt = formato.lower().strip()
    if fmt not in ("pdf", "xlsx", "docx"):
        raise HTTPException(status_code=400, detail="Formato no soportado")

    if titulo is None:
        titulo = ""
    titulo = str(titulo)

    if columnas is None:
        columnas = []
    if not isinstance(columnas, (list, tuple)):
        raise HTTPException(status_code=400, detail="Columnas debe ser una lista")

    if datos is None:
        datos = []
    if not hasattr(datos, '__iter__'):
        raise HTTPException(status_code=400, detail="Datos debe ser iterable de filas")

    rows = _normalize_rows(list(columnas), list(datos))

    try:
        if fmt == "pdf":
            pdf = FPDF(orientation='L')
            pdf.add_page()
            pdf.set_font("helvetica", "B", 16)

            pdf.set_text_color(30, 58, 138)
            pdf.cell(0, 10, titulo, ln=True, align="C")
            pdf.ln(8)

            pdf.set_font("helvetica", "B", 10)
            pdf.set_fill_color(30, 58, 138)
            pdf.set_text_color(255, 255, 255)

            ancho_col = 270 / max(1, len(columnas))
            if columnas:
                for col in columnas:
                    pdf.cell(ancho_col, 10, str(col), border=1, align="C", fill=True)
            else:
                # Si no hay columnas, dibujamos una sola celda con el título
                pdf.cell(0, 10, "Datos", border=1, align="C", fill=True)
            pdf.ln()

            pdf.set_font("helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            for fila in rows:
                for item in fila:
                    texto = str(item)
                    if len(texto) > 60:
                        texto = texto[:57] + "..."
                    pdf.cell(ancho_col, 8, texto, border=1, align="C")
                pdf.ln()

            raw = pdf.output(dest='S')
            if isinstance(raw, (bytes, bytearray)):
                pdf_bytes = bytes(raw)
            else:
                pdf_bytes = str(raw).encode('latin-1', errors='replace')
            return StreamingResponse(io.BytesIO(pdf_bytes), media_type="application/pdf")

        if fmt == "xlsx":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = titulo[:31] or "Reporte"

            if columnas:
                ws.append([str(c) for c in columnas])
            for fila in rows:
                ws.append([str(v) for v in fila])

            # Ajuste básico de ancho de columnas
            for col in ws.columns:
                try:
                    max_length = max((len(str(cell.value)) for cell in col if cell.value is not None), default=0)
                    column = col[0].column_letter
                    ws.column_dimensions[column].width = max(10, max_length + 2)
                except Exception:
                    pass

            stream = io.BytesIO()
            wb.save(stream)
            stream.seek(0)
            return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        if fmt == "docx":
            doc = Document()
            if titulo:
                doc.add_heading(titulo, 0)

            cols = len(columnas) if columnas else (len(rows[0]) if rows else 1)
            tabla = doc.add_table(rows=1, cols=max(1, cols))
            tabla.style = 'Table Grid'

            if columnas:
                hdr_cells = tabla.rows[0].cells
                for i, col in enumerate(columnas):
                    if i < len(hdr_cells):
                        hdr_cells[i].text = str(col)
            else:
                tabla.rows[0].cells[0].text = ""

            for fila in rows:
                row_cells = tabla.add_row().cells
                for i, item in enumerate(fila):
                    if i < len(row_cells):
                        row_cells[i].text = str(item)

            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return StreamingResponse(stream, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando archivo: {e}")

    # Fallback (shouldn't alcanzarse)
    raise HTTPException(status_code=400, detail="Formato no soportado")
